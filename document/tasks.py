from __future__ import absolute_import

import pycorrector, io, docx, subprocess, os, base64
from django.utils import timezone
import logging

from docx.shared import RGBColor

from NLP.celery import app
from NLP.settings import BASE_DIR
from document.models import Document
from utils.libs import baidu_correct, contain_zh


def do_docx(file_bytes, file_name, username, process_method):
    s = io.BytesIO(file_bytes)
    doc = docx.Document(s)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if contain_zh(run.text):
                if process_method == 'kenlm':
                    correct_sent = pycorrector.correct(run.text)[1]
                else:
                    correct_sent = baidu_correct(run.text)
                if correct_sent:
                    run.font.color.rgb = RGBColor(255, 0, 0)
        # 调用api,获取返回的数据
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if contain_zh(run.text):
                            if process_method == 'kenlm':
                                correct_sent = pycorrector.correct(run.text)[1]
                            else:
                                correct_sent = baidu_correct(run.text)
                                # 调用api,获取返回的数据
                            if correct_sent:
                                run.font.color.rgb = RGBColor(255, 0, 0)
    dir = os.path.join(BASE_DIR, 'utils/{}'.format(username))
    if not os.path.exists(dir):
        os.makedirs(dir)
    doc.save(os.path.join(dir, file_name))
    return dir


def do_csv():
    pass


def do_txt(file_bytes, file_name, username, process_method):
    dir = os.path.join(BASE_DIR, 'utils/{}'.format(username))
    # 1 分段，段超过511字节，结巴分词
    file_str = file_bytes.decode()
    lines = file_str.splitlines()
    with open(os.path.join(dir, file_name), 'w') as f:
        for l in lines:
            if l and contain_zh(l):
                if process_method != 'kenlm':
                    query = baidu_correct(l)
                else:
                    query = pycorrector.correct(l)[0]
            else:
                query = l
            f.write(query)
    return dir


@app.task(bind=True)
def process_file(self, file_type, file_bytes, file_name, username, doc_id, process_method='kenlm'):
    doc = Document.objects.filter(id=doc_id).first()
    file_bytes = base64.b64decode(file_bytes.encode())
    if doc:
        try:
            if file_type == 'doc':
                output = subprocess.check_output(
                    ["soffice", "--headless", "--invisible", "--convert-to", "docx", f"{file_name}", "--outdir", ""])
                # path = output.decode().splitlines()[1].split('Overwriting:')[1].strip()
                os.remove(file_name)
                file_name = f'{file_name[:-4]}.docx'
                with open(file_name, 'rb') as f:
                    dir = do_docx(f.read(), file_name, username, process_method)
                os.remove(file_name)
            elif file_type == 'docx':
                dir = do_docx(file_bytes, file_name, username, process_method)
            elif file_type == 'csv':
                dir = do_csv()
            elif file_type == 'txt' or file_type == 'text':
                dir = do_txt(file_bytes, file_name, username, process_method)
            else:
                raise ValueError('不被支持的类型!')
            doc.status = 3
            doc.success_url = os.path.join(dir, file_name)
            doc.st = timezone.now()
        except Exception as e:
            logging.error(repr(e))
            doc.status = 2
        doc.save()
