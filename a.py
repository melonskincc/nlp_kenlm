import re

import docx, jieba
import pycorrector
import kenlm

# m = kenlm.LanguageModel('a.arpa')
# print(m.score('', bos = True, eos = True))

# model = pycorrector.Corrector(language_model_path='a.arpa')
# corrected_sent, detail = model.correct('少先队员因该为老人让坐')
# print(corrected_sent, detail)
# s = '2012年 醉新 普通话 等极 考试（规则、练习材料、试题）'
# err = pycorrector.correct(s)
# print(err)
# pycorrector.language_model_path('/root/.pycorrector/datasets/it.arpa')
# x = model.correct(s)
# print(x)
# s = '石瘤'
# err = pycorrector.correct(s)
# print(err)
# pycorrector.language_model_path('/root/.pycorrector/datasets/it.arpa')
# x = model.correct(s)
# print(x)
# s = '少先队员因该为老人让坐'
# err = pycorrector.correct(s)
# print(err)
# pycorrector.language_model_path('/root/.pycorrector/datasets/it.arpa')
# x = model.correct(s)
# print(x)
# y = pycorrector.correct(s)
# print(y)
# error_sentence_1 = '我的喉咙发炎了要买点阿莫细林吃'
# error_sentence_1 = '少先队员因该为老人让坐'
# error_sentence_1 = '这是醉新到板本了'
# s = jieba.cut(error_sentence_1)
# for x in s:
#     print(x)
#     correct_sent = pycorrector.correct(x)
#     print(correct_sent)
# pycorrector.enable_char_error(enable=False)
# correct_sent = pycorrector.correct(error_sentence_1)
# print(correct_sent)
import subprocess
output = subprocess.check_output(["soffice","--headless","--invisible","--convert-to","docx","普通话等级考试资料(规则、材料、试.....doc","--outdir","."])
print(output.decode())
# fn = './普通话等级考试资料(规则、材料、试.....docx'
# doc = docx.Document(fn)
# pycorrector.set_language_model_path('a.arpa')
# print(pycorrector.word_frequency('等级'))
# count = 0
# for paragraph in doc.paragraphs:
#     print(paragraph.text)
#     correct_sent = pycorrector.correct(paragraph.text)
#     paragraph.text = correct_sent[0]
#     count += 1
#     if count == 20:
#         break
    # print("处理后:", correct_sent)
    # for x in jieba.lcut(re.sub(r'[^\u4e00-\u9fa5]', ' ', correct_sent[0])):
    #     print(x)
#
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
#             correct_sent = pycorrector.correct(cell.text)
#             print("处理后:", correct_sent)
#

# doc.save('test.docx')
